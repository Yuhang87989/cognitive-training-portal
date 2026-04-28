from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, border_color="000000", border_size="4"):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:left w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

def add_heading(doc, text, level):
    """添加标题"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.name = 'SimHei'
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    return heading

def add_paragraph(doc, text, font_size=11):
    """添加正文段落"""
    para = doc.add_paragraph(text)
    for run in para.runs:
        run.font.name = 'SimSun'
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    return para

def add_code_block(doc, text):
    """添加代码块"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return para

def create_table(doc, headers, rows, col_widths=None):
    """创建表格"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.text = header
        set_cell_shading(cell, "D9D9D9")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.name = 'SimSun'
                run.font.size = Pt(10)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        set_cell_border(cell)
    
    # 数据行
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_text in enumerate(row_data):
            cell = row_cells[col_idx]
            cell.text = str(cell_text)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'SimSun'
                    run.font.size = Pt(9)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            set_cell_border(cell)
    
    return table

# 创建文档
doc = Document()

# 设置页面边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ===== 封面页 =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('认知训练门户——本周有效沟通与改进历程报告')
run.font.name = 'SimHei'
run.font.bold = True
run.font.size = Pt(22)
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('从V20到V140模块化重构的完整历程')
run.font.name = 'SimSun'
run.font.size = Pt(14)
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

doc.add_paragraph()
doc.add_paragraph()

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run('2026年4月22日-4月28日')
run.font.name = 'SimSun'
run.font.size = Pt(12)
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

doc.add_paragraph()
# 横线
line_para = doc.add_paragraph('_' * 60)
line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# ===== 一、项目概述 =====
add_heading(doc, '一、项目概述', 1)

add_heading(doc, '项目基本信息', 2)
info_items = [
    '项目名称：青少年认知训练门户',
    '目标用户：12-16岁青少年',
    '核心功能：12大模块全面覆盖'
]
for item in info_items:
    add_paragraph(doc, item)

add_heading(doc, '12大核心模块', 2)
modules = [
    'AI数字分身、AI精准练、认知地图、学习计划',
    '母题训练、学霸方法、思维训练',
    '播客课堂、视频课堂',
    '训练游戏（23认知+8娱乐=31个）',
    'DeepSeek AI、设置'
]
for m in modules:
    add_paragraph(doc, m)

add_heading(doc, '双平台部署', 2)
platforms = [
    'GitHub Pages：yuhang87989.github.io/cognitive-training-portal/',
    '扣子平台：6mz7txgx3f.coze.site'
]
for p in platforms:
    add_paragraph(doc, p)

add_heading(doc, '默认用户', 2)
add_paragraph(doc, '用户名：邱宇菲（初一）')
add_paragraph(doc, '难度级别：1')
add_paragraph(doc, '初始积分：1142')

# ===== 二、本周版本演进历程 =====
add_heading(doc, '二、本周版本演进历程', 1)

version_headers = ['日期', '版本', '主要变更', '核心问题']
version_rows = [
    ['4月22日', 'V20', '首次12模块布局，8大训练游戏，底部导航', '基础框架搭建'],
    ['4月22日', 'V21', '多用户登录，数据同步，用户切换', '登录页面功能'],
    ['4月22日', 'V22', '首页布局优化，浅蓝渐变背景', '视觉优化'],
    ['4月22日', 'V24', '修复创建新账号按钮', '按钮不可用'],
    ['4月22日', 'V26', '12大模块全部就位，底部5标签导航', '模块完整性'],
    ['4月22日', 'V27', '61个模块检测通过，42个函数', '功能验证'],
    ['4月22日', 'V29', '修复9项（视觉搜索、找不同、AI出题等）', '游戏逻辑错误'],
    ['4月22日', 'V30', '数据迁移，登录异常处理，清除数据', '数据迁移'],
    ['4月22日', 'V32', '登录页面不显示', 'V30到V32回归问题'],
    ['4月22日', 'V33', '绕过扣子部署系统Bug', '扣子平台兼容'],
    ['4月22日', 'V37', '修复JS语法错误，页面空白', '稳定版本'],
    ['4月23日', 'V50-V56', '多版本迭代，游戏全屏、UI优化', '持续迭代'],
    ['4月26日', 'V72-V81', '模块点击无响应、缺失函数、HTML残留修复', '功能补全'],
    ['4月26日', 'V95', 'topics语法错误，清理重复函数', '代码质量'],
    ['4月27日', 'V117-V120', '多用户完善、默认用户机制、设置面板', '用户系统'],
    ['4月27日', 'V126-V128', '播客音频URL、视频上传、签名URL', '媒体功能'],
    ['4月28日', 'V139', '模块化升级，结构化检查', '架构转型'],
    ['4月28日', 'V140', '完整模块化重构：单文件到26个JS模块+CSS', '架构重构'],
]
create_table(doc, version_headers, version_rows)

# ===== 三、关键沟通节点分析 =====
add_heading(doc, '三、关键沟通节点分析', 1)

comm_headers = ['时间', '用户反馈', 'AI理解偏差', '核心问题', '改进方向']
comm_rows = [
    ['4/28 19:38', '要求上传部署文件并部署', '正确执行', '双平台需同步', '建立双平台部署流程'],
    ['4/28 20:10', '游戏数量应为23个不是16个', '文案错误', '数字不准确', '全面核对实际数量'],
    ['4/28 20:21', '两个平台部署文件不一样', '文件不同步', '一方能用一方不能', '必须双平台同步修改'],
    ['4/28 20:22', '图标没改、返回按钮没有', '视觉细节遗漏', '用户体验', '细节检查清单化'],
    ['4/28 20:27', '双轨方案修改文件不一样', '子任务只改一个平台', '部署不同步', '改完必须双平台验证'],
    ['4/28 21:44', 'GitHub部署不能用了', '全面检查发现根因', '单文件不可维护', '必须模块化'],
    ['4/28 21:55', '之前有模块化文件，完善再部署', '用户指出正确方向', 'AI走弯路', '听用户方向立即执行'],
    ['4/28 21:56', '不要动设计师思路', '过度修改', '尊重原有设计', '完善而非重写'],
    ['4/28 21:59', '扣子平台是否自己修改了？', '关键发现', '扣子支持多文件', '打破错误认知'],
    ['4/28 22:01', '稳定版就是模块化，完善就好', 'AI犹豫不决', '用户坚定方向', '坚定执行用户意图'],
    ['4/28 22:10', '必须做结构化模块化才能维护', '核心决策', '反复修补不可行', '立即执行模块化'],
    ['4/28 22:13', '反复强调"执行整体模块化修改"', '用户不满', 'AI理解偏慢', '直接行动不反复确认'],
    ['4/28 22:17', '你觉得太慢才来执行', '效率不够', '明确要求', '主Agent直接接手'],
    ['4/28 22:38', '记住部署步骤，不然变死循环', '信息丢失', '防止重复踩坑', '持久化记录关键信息'],
    ['全周', '花200多元积分完善系统', '成本意识', '效率问题', '每次修改要高效准确'],
]
create_table(doc, comm_headers, comm_rows)

# ===== 四、沟通效率反思 =====
add_heading(doc, '四、沟通效率反思', 1)

reflections = [
    ('1. 方向确认后立即执行', 
     '用户多次强调模块化方向，AI初期理解偏差走了弯路。当用户明确方向后应立即执行，不要反复确认"是不是这个意思"，这会让用户非常焦虑。反复确认看似谨慎，实则降低了效率。'),
    ('2. 单文件修补是死循环', 
     '780KB单文件改一处坏三处，维护成本指数增长。本周最重要的教训——模块化才是可维护的正确方向，永远不要回到单文件模式。单文件模式看似方便，长期维护代价极高。'),
    ('3. 尊重原有设计', 
     '"不要动设计师思路，不完善就完善，有错误就改正"——不重写、不改变设计意图，只做补全和修正。过度修改反而引入新问题。保持设计一致性比推倒重来更有效率。'),
    ('4. 双平台必须同步', 
     '修改必须同步到GitHub Pages和扣子平台，否则出现"一个能用一个不能用"的问题。子任务修改时容易只改一个平台，必须建立同步检查机制。'),
    ('5. 记录防止死循环', 
     '部署步骤、文件位置、平台差异必须持久化记录到MEMORY.md，否则每次都从头摸索，浪费大量时间和积分。用户明确说"不然又要变成死循环"。'),
    ('6. 用户成本意识', 
     '用户花了200多元积分购买服务，时间成本更不可估量。每次修改要高效准确，减少无效迭代，避免"修了A坏了B"的情况。提高一次成功率，减少试错成本。'),
]

for title_text, content_text in reflections:
    add_heading(doc, title_text, 2)
    add_paragraph(doc, content_text)

# ===== 五、技术改进总结 =====
add_heading(doc, '五、技术改进总结', 1)

tech_headers = ['改进领域', '改进前', '改进后', '效果']
tech_rows = [
    ['整体架构', '单文件780KB内联', '26个JS+1个CSS+1个HTML(39KB)', '改一个模块只动一个文件'],
    ['CSS管理', '内联style属性', '独立style.css(26KB)', '样式可独立修改'],
    ['游戏文案', '"8种"/"16大"', '23认知+8娱乐=31个', '准确反映实际'],
    ['返回按钮', '图标不明显', '文字+渐变背景', '不迷路'],
    ['Service Worker', '旧缓存+错误路径', 'V140网络优先+路径修正', '加载最新版'],
    ['图标路径', '引用不存在文件', '修正为实际文件', 'PWA正常'],
    ['播客播放', '占位URL', 'coze CDN真实地址', '可播放'],
    ['双平台', '文件不同步', '统一模块化源文件', '一次改双平台生效'],
    ['函数管理', '594函数单文件', '404具名函数26模块', '独立可维护'],
    ['扣子认知', '认为不支持多文件(错)', '确认js/css目录可用', '模块化部署打通'],
]
create_table(doc, tech_headers, tech_rows)

# ===== 六、模块化架构图 =====
add_heading(doc, '六、模块化架构图', 1)

add_heading(doc, '文件结构树', 2)
code_content = """cognitive-training-portal/
├── index.html (39KB) - HTML结构+脚本加载器
├── css/style.css (26KB) - 主样式表
├── js/config.js - 全局配置
├── js/ctm.js - 模块管理器
├── js/audio.js - 音效系统
├── js/storage.js - 存储模块
├── js/utils.js - 工具函数
├── js/user.js - 用户系统
├── js/data/week-plans.js - 训练计划(21KB)
├── js/data/topics.js - 母题题库(59KB)
├── js/data/podcasts.js - 播客数据
├── js/data/videos.js - 视频数据
├── js/data/games-config.js - 游戏配置(10KB)
├── js/modules/practice.js - 母题训练
├── js/modules/map.js - 认知地图
├── js/modules/plan.js - 学习计划
├── js/modules/topics.js - 科目选择
├── js/modules/method.js - 学霸方法
├── js/modules/thinking.js - 思维训练
├── js/modules/podcast.js - 播客课堂(48KB)
├── js/modules/video.js - 视频课堂
├── js/modules/player.js - 播放器(38KB)
├── js/modules/games.js - 31个游戏(297KB)
├── js/modules/deepseek.js - DeepSeek AI(38KB)
├── js/modules/wrongbook.js - 错题本
├── js/modules/pomodoro.js - 番茄钟
├── js/modules/ui.js - UI主控(60KB)
├── js/modules/ai.js - AI辅助
├── manifest.json / service-worker.js
└── favicon.ico / icon-192.png / apple-touch-icon.png"""
add_code_block(doc, code_content)

add_heading(doc, 'JS加载顺序', 2)
load_order = "config->ctm->audio->storage->utils->user->data/*->practice->map->plan->topics->method->thinking->podcast->video->player->games->deepseek->wrongbook->pomodoro->ui->ai"
add_paragraph(doc, load_order)

# ===== 七、双平台部署对比 =====
add_heading(doc, '七、双平台部署对比', 1)

deploy_headers = ['项目', 'GitHub Pages', '扣子平台']
deploy_rows = [
    ['部署方式', 'git push自动', 'coze code消息+deploy'],
    ['地址', 'yuhang87989.github.io/cognitive-training-portal/', '6mz7txgx3f.coze.site'],
    ['文件路径', '相对路径', '相对路径(支持多文件)'],
    ['PWA', '完整支持', 'SPA路由可能受限'],
    ['注入', '无', '自动注入coze-init.js'],
    ['当前版本', 'V140模块化', 'V140模块化'],
    ['项目ID', 'GitHub: Yuhang87989', '7627567739896381459'],
]
create_table(doc, deploy_headers, deploy_rows)

# ===== 八、版本里程碑统计 =====
add_heading(doc, '八、版本里程碑统计', 1)

milestone_headers = ['里程碑', '版本范围', '关键成果']
milestone_rows = [
    ['基础搭建', 'V20-V27', '12大模块框架、8大游戏、多用户登录'],
    ['功能完善', 'V28-V38', '数据迁移、登录修复、扣子兼容'],
    ['持续迭代', 'V50-V81', '游戏全屏、UI优化、函数补全、代码清理'],
    ['用户系统', 'V95-V120', '多用户完善、设置面板、数据保护'],
    ['媒体功能', 'V126-V128', '播客音频、视频上传、签名URL播放'],
    ['架构转型', 'V139-V140', '模块化重构，单文件到26个JS模块'],
]
create_table(doc, milestone_headers, milestone_rows)

# ===== 九、下一步计划 =====
add_heading(doc, '九、下一步计划', 1)

plans = [
    '验证12大模块全链路功能完整性',
    '错题本拍照上传+AI解说功能',
    'APK重新打包（V140模块化版）',
    '全局变量优化（CTM命名空间封装）',
    'Week1-7训练计划数据完善',
    '持续模块化优化，坚决不回退单文件模式'
]
for i, plan in enumerate(plans, 1):
    add_paragraph(doc, f'{i}. {plan}')

# 保存文档 - 使用绝对路径和ASCII文件名
base_dir = '/app/data/所有对话/主对话/cognitive-training-portal'
output_path = os.path.join(base_dir, 'weekly_report.docx')
doc.save(output_path)
print(f'文档已生成: {output_path}')
